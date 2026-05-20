"""Extraction de texte universelle pour documents.

Architecture extensible : chaque format a un extracteur dédié.
V2 ajoutera MP3/MP4 (transcription audio/vidéo).
"""
from __future__ import annotations

import hashlib
import io
import re
from pathlib import Path

import structlog

log = structlog.get_logger()

# Extensions supportées actuellement
_TEXT_EXT = frozenset({".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm", ".log", ".rst"})
_DOC_EXT = frozenset({".docx", ".doc"})
_PDF_EXT = frozenset({".pdf"})
_IMAGE_EXT = frozenset({".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".gif", ".webp"})

# Future V2
_AUDIO_EXT = frozenset({".mp3", ".wav", ".m4a", ".ogg", ".flac"})
_VIDEO_EXT = frozenset({".mp4", ".avi", ".mov", ".mkv", ".wmv"})

_MAX_CHUNK_SIZE = 4000  # tokens approximatifs


def _chunk_text(text: str, chunk_size: int = _MAX_CHUNK_SIZE) -> list[str]:
    """Découpe un texte long en chunks avec chevauchement.

    Utilise les sauts de ligne comme frontières naturelles quand possible.
    """
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    overlap = chunk_size // 10  # 10% chevauchement
    start = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))
        # Chercher un saut de ligne proche pour couper proprement
        if end < len(text):
            nl_pos = text.rfind("\n", end - chunk_size // 5, end)
            if nl_pos != -1:
                end = nl_pos + 1
        chunks.append(text[start:end].strip())
        start = end - overlap

    return [c for c in chunks if c]


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        log.warning("document_extract.pdf_lib_missing")
        return ""

    try:
        reader = PdfReader(str(path))
        pages: list[str] = []
        for i, page in enumerate(reader.pages):
            txt = page.extract_text() or ""
            if txt.strip():
                pages.append(f"--- Page {i + 1} ---\n{txt.strip()}")
        return "\n\n".join(pages)
    except Exception as e:
        log.warning("document_extract.pdf_failed", path=str(path), error=str(e))
        return ""


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        log.warning("document_extract.docx_lib_missing")
        return ""

    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        log.warning("document_extract.docx_failed", path=str(path), error=str(e))
        return ""


def _extract_image(path: Path) -> str:
    try:
        from PIL import Image
    except ImportError:
        log.warning("document_extract.pil_missing")
        return ""

    try:
        # Vérifier que tesseract est dispo
        import pytesseract
        img = Image.open(path)
        text = pytesseract.image_to_string(img)
        return text.strip()
    except Exception as e:
        log.warning("document_extract.image_ocr_failed", path=str(path), error=str(e))
        return ""


_EXTRACTORS: dict[str, callable] = {}
for ext in _TEXT_EXT:
    _EXTRACTORS[ext] = _extract_txt
for ext in _PDF_EXT:
    _EXTRACTORS[ext] = _extract_pdf
for ext in _DOC_EXT:
    _EXTRACTORS[ext] = _extract_docx
for ext in _IMAGE_EXT:
    _EXTRACTORS[ext] = _extract_image


def _normalize_ext(ext: str) -> str:
    return ext.lower().lstrip(".")


def _ext_to_dot(ext: str) -> str:
    e = ext.lower()
    return f".{e}" if not e.startswith(".") else e


def extract_text(path: Path | str) -> str:
    """Extrait le texte brut d'un fichier selon son extension.

    Retourne "" si le format est inconnu ou si l'extraction échoue.
    """
    p = Path(path)
    ext = p.suffix.lower()

    if not p.exists():
        log.warning("document_extract.not_found", path=str(p))
        return ""

    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        log.info("document_extract.unsupported_format", ext=ext, path=str(p))
        return ""

    return extractor(p)


def extract_text_bytes(data: bytes, filename: str) -> str:
    """Extrait le texte depuis des bytes en mémoire (pièces jointes emails)."""
    ext = Path(filename).suffix.lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        log.info("document_extract.unsupported_format", ext=ext, filename=filename)
        return ""

    # Les extracteurs travaillent sur Path ; on écrit temporairement
    if ext in _TEXT_EXT:
        return data.decode("utf-8", errors="replace")

    try:
        from pypdf import PdfReader
        if ext in _PDF_EXT:
            reader = PdfReader(io.BytesIO(data))
            pages = []
            for i, page in enumerate(reader.pages):
                txt = page.extract_text() or ""
                if txt.strip():
                    pages.append(f"--- Page {i + 1} ---\n{txt.strip()}")
            return "\n\n".join(pages)
    except Exception as e:
        log.warning("document_extract.pdf_bytes_failed", filename=filename, error=str(e))

    try:
        from docx import Document
        if ext in _DOC_EXT:
            doc = Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
    except Exception as e:
        log.warning("document_extract.docx_bytes_failed", filename=filename, error=str(e))

    try:
        from PIL import Image
        import pytesseract
        if ext in _IMAGE_EXT:
            img = Image.open(io.BytesIO(data))
            return pytesseract.image_to_string(img).strip()
    except Exception as e:
        log.warning("document_extract.image_bytes_failed", filename=filename, error=str(e))

    return ""


def chunk_text(text: str) -> list[str]:
    """Découpe intelligente du texte pour ingestion par morceaux."""
    return _chunk_text(text)


def content_hash(text: str) -> str:
    """Hash MD5 du contenu pour détection de doublons."""
    return hashlib.md5(text.encode("utf-8"), usedforsecurity=False).hexdigest()


def is_supported(filename: str) -> bool:
    """Vérifie si l'extension est supportée par l'extracteur actuel."""
    ext = Path(filename).suffix.lower()
    return ext in _EXTRACTORS


def list_supported_extensions() -> list[str]:
    """Retourne la liste des extensions supportées (sans le point)."""
    exts = set()
    for e in _EXTRACTORS:
        exts.add(e.lstrip("."))
    return sorted(exts)
